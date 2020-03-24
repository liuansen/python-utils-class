# -*- coding:utf-8 -*-
# author：Anson
from __future__ import unicode_literals

import os
import sys
import re
from datetime import date, datetime, timedelta

from docx import Document
import xlwt

from settings import MD_PATH, SITE_1, SITE_2, CELL


reload(sys)
sys.setdefaultencoding('utf-8')


def get_file_path(path, week_of, table1, table2, first_date, today, worksheet, site_1, site_2,
                  first_date_of, today_of):
    style = xlwt.XFStyle()
    bl = xlwt.Borders()
    bl.left = xlwt.Borders.THIN
    bl.right = xlwt.Borders.THIN
    bl.top = xlwt.Borders.THIN
    bl.bottom = xlwt.Borders.THIN
    al = xlwt.Alignment()
    al.horz = 0x02  # 设置水平居中
    al.vert = 0x01  # 设置垂直居中
    style.alignment = al
    style.borders = bl

    nums = 0
    file_date = date.today().strftime('%Y-%m')
    for filename in os.listdir(path):
        file_path = os.path.join(path, filename)
        group_name = re.findall(r'.*2019-08-(.*)..*', filename)[0][0:-2]
        fd = filename[:7]
        md = file_path[-2:]
        if md == 'md':
            if fd == file_date:
                with open(file_path) as f:
                    lines = f.readlines()
                lines = [i.strip('-').strip() for i in lines]
                if len(lines) == 0:
                    first_index = 0
                else:
                    for key, value in enumerate(lines):
                        if value == week_of:
                            first_index = key
                        else:
                            first_index = 0
                k = 0
                line_list = []
                index = 0
                while k < len(lines):
                    if lines[k] == week_of:
                        index += 1
                        first_index = k
                        line_list.append(lines[k])
                    else:
                        if k > first_index:
                            if lines[k][:1] == '#':
                                break
                            else:
                                line_list.append(lines[k])
                    k += 1
                line = [i.strip('#').strip() for i in line_list]
                d = 0
                trade_today = False
                yearst_today = False
                s1 = ''
                s2 = ''
                sor_index = 0
                while d < len(line):
                    if line[d].strip()[:1] == '*':
                        if sor_index != 0:
                            worksheet.write(site_1, 2, s1, style)
                            worksheet.write(site_2, 2, s2, style)
                            s1 = ''
                            s2 = ''
                        yearst_today = False
                        nums += 1
                        site_1 += 1
                        site_2 += 1
                        name = line[d].strip('*').strip()
                        worksheet.write(site_1, 1, str(nums), style)
                        worksheet.write(site_1, 3, first_date, style)
                        worksheet.write(site_1, 4, today, style)
                        worksheet.write(site_1, 5, name, style)
                        worksheet.write(site_2, 1, str(nums), style)
                        worksheet.write(site_2, 3, first_date_of, style)
                        worksheet.write(site_2, 4, today_of, style)
                        worksheet.write(site_2, 5, name, style)
                        table1.rows[nums].cells[0].add_paragraph(str(nums))
                        table1.rows[nums].cells[2].add_paragraph(first_date)
                        table1.rows[nums].cells[3].add_paragraph(today)
                        table1.rows[nums].cells[4].add_paragraph(name)
                        table1.rows[nums].cells[5].add_paragraph(group_name)
                        table2.rows[nums].cells[0].add_paragraph(str(nums))
                        table2.rows[nums].cells[2].add_paragraph(first_date_of)
                        table2.rows[nums].cells[3].add_paragraph(today_of)
                        table2.rows[nums].cells[4].add_paragraph(name)
                        table2.rows[nums].cells[5].add_paragraph(group_name)
                        d += 1
                        sor_index += 1
                    if line[d] == '本周工作':
                        trade_today = True
                        d += 1
                    if (line[d].strip()[1:2] == '.' or line[d].strip()[1:2] == ')') and trade_today:
                        # 本周工作内容
                        table1.rows[nums].cells[1].add_paragraph(line[d])
                        s1 = s1 + ' ' + line[d]
                    if line[d] == '下周工作' or line[d] == '下周计划':
                        trade_today = False
                        yearst_today = True
                        d += 1
                    if (line[d].strip()[1:2] == '.' or line[d].strip()[1:2] == ')') and yearst_today:
                        # 下周工作内容
                        table2.rows[nums].cells[1].add_paragraph(line[d])
                        s2 = s2 + ' ' + line[d]
                    d += 1
                worksheet.write(site_1, 2, s1, style)
                worksheet.write(site_2, 2, s2, style)


def get_week_of_month(year, month, day):
    """
    获取指定的某天是某个月中的第几周
    周一作为一周的开始
    """
    end = int(datetime(year, month, day).strftime("%W"))
    begin = int(datetime(year, month, 1).strftime("%W"))
    star_date = end - begin + 1
    if star_date == 1:
        week_of = '# 第一周'
    elif star_date == 2:
        week_of = '# 第二周'
    elif star_date == 3:
        week_of = '# 第三周'
    elif star_date == 4:
        week_of = '# 第四周'
    elif star_date == 5:
        week_of = '# 第五周'
    else:
        week_of = '# 第六周'
    return week_of


def create_table_one_cell(document, content):
    """创建单行列表"""
    create_table = document.add_table(rows=1, cols=1, style='Table Grid')
    create_table.rows[0].cells[0].add_paragraph(content)


def create_table_more_cell(document, rows, cols, contents):
    """创建多行多列的列表"""
    create_table = document.add_table(rows=rows, cols=cols, style='Table Grid')
    index = 0
    for content in contents:
        for key, value in enumerate(content):
            create_table.rows[index].cells[key].add_paragraph(value)
        index += 1


def create_fixed_cell(document, first_date, end_date):
    """表前半部分固定内容"""
    create_table_one_cell(document, '项目基本情况')
    create_table_more_cell(document, 2, 2, [['项目名称', '厦开项目组'], ['客户名称', '中国建设银行厦门开发中心']])
    create_table_more_cell(document, 3, 6, [['客户负责人', '李晓敦', '电话', '', 'Email', ''],
                                            ['(必填)', '闫立志', '电话', '', 'Email', ''],
                                            ['', '', '电话', '', 'Email', '']])
    create_table_more_cell(document, 4, 2, [['开始日期', first_date], ['项目经理', '赖志勇'],
                                            ['项目组成员', '柳惠阳、许华语、郭健超、何卧岩、郑炜、黄惠章、朱俊龙、李稳定、'
                                                      '黄建鸣、陈浩1、叶晟君、张叶桃、陈晓衍、曾国荣、肖凯、刘安森、'
                                                      '林秋霞、姜渊、肖金平、周丽荣、钟晓杰、黄祯鸿、李志阳、刘程川、'
                                                      '张俊钦、邓松进、林丹丹、姜琪、钟高镇、方若琳、、谢源鑫、罗庭颖、'
                                                      '魏治邦、白艺伟、付敏、肖金龙、颜炳煜、庄华琼、董凯华、黄忠强、'
                                                      '徐鸿能、江养根、何龙伙、肖丽琴、罗万春、曾林华、、张一浓、郭吉、、'
                                                      '吴招辉、林泉、、苏雪梅、张祖琦、、陈浩'],
                                            ['项目描述', '']])
    create_table_one_cell(document, '计划关键时间点（必填）')
    create_table_more_cell(document, 6, 4, [['关键时间点', '预计完成时间', '关键时间点', '预计完成时间'],
                                            ['1、需求分析', '', '6、技术测试(单元测试）', ''],
                                            ['2、技术方案(项目实施方案)', '', '7、业务测试(集成测试)', ''],
                                            ['3、概要设计', '', '8、上线时间', ''],
                                            ['4、详细设计', '', '9、后期维护', ''],
                                            ['5、编码', '', '10、结项', '']])
    create_table_one_cell(document, '实际关键时间点（必填）')
    create_table_more_cell(document, 6, 4, [['关键时间点', '实际完成时间', '关键时间点', '实际完成时间'],
                                            ['1、需求分析', '', '6、技术测试(单元测试）', ''],
                                            ['2、技术方案(项目实施方案)', '', '7、业务测试(集成测试)', ''],
                                            ['3、概要设计', '', '8、上线时间', ''],
                                            ['4、详细设计', '', '9、后期维护', ''],
                                            ['5、编码', '', '10、结项', '']])
    create_table_one_cell(document, '人力资源状况（包括人员的入职、离职;入场、离场、休假、请假等情况）.'
                                    '时间以到达、离开现场为准')
    create_table_one_cell(document, '预计新增资源（必填）')
    create_table_more_cell(document, 4, 6, [['姓名', '', '预计到场时间', '', '任务描述', ''],
                                            ['姓名', '', '预计到场时间', '', '任务描述', ''],
                                            ['姓名', '', '预计到场时间', '', '任务描述', ''],
                                            ['姓名', '', '预计到场时间', '', '任务描述', '']])
    create_table_one_cell(document, '预计撤离资源（必填）')
    create_table_more_cell(document, 3, 6, [['姓名', '', '预计离场时间', '', '撤离原因', ''],
                                            ['姓名', '', '预计离场时间', '', '撤离原因', ''],
                                            ['姓名', '', '预计离场时间', '', '撤离原因', '']])
    create_table_one_cell(document, '本周人员变动情况（必填）')
    create_table_more_cell(document, 5, 4, [['序号', '到场人员姓名', '到场时间', '备注'],
                                            ['1', '', '', ''], ['2', '', '', ''],
                                            ['3', '', '', ''], ['4', '', '', '']])
    create_table_more_cell(document, 5, 4, [['序号', '离场人员姓名', '离场时间', '备注'],
                                            ['1', '', '', ''], ['2', '', '', ''],
                                            ['3', '', '', ''], ['4', '', '', '']])
    create_table_one_cell(document, '本周项目情况')
    create_table_one_cell(document, '项目所处阶段（必填）')
    create_table_more_cell(document, 2, 5, [['1、需求分析', '2、概要设计', '3、详细设计', '4、编码', '5、技术测试'],
                                            ['6、业务测试', '7、试运行	', '8、部分上线', '9、整体完工',	'10、后期维护']])
    create_table_one_cell(document, '项目经理自评（必填）')
    create_table_more_cell(document, 5, 2, [['是否完成以下事项', '未完成的理由及说明'],
                                            ['是否组织周例会会议纪要？ 【□是   □否】', ''],
                                            ['本周工作是否按计划完成？【□是   □否】', ''],
                                            ['是否跟客户项目负责人汇报本周工作？【□是 □否】', ''],
                                            ['下周计划安排是否与项目成员落实？【□是 □否】	', '']])
    create_table_one_cell(document, '需求变更情况（必填）')
    create_table_more_cell(document, 3, 2, [['需求变更描述', '对后续的影响'], ['无', ''], ['', '']])
    create_table_one_cell(document, '方案变更情况（必填）')
    create_table_more_cell(document, 3, 2, [['方案变更描述', '对后续的影响'], ['', ''], ['', '']])
    create_table_one_cell(document, '项目计划变更情况（必填）')
    create_table_more_cell(document, 3, 2, [['项目计划变更描述', '对后续的影响'], ['', ''], ['', '']])
    create_table_one_cell(document, '本周未完成的任务情况（必填）')
    create_table_more_cell(document, 4, 3, [['未完成的任务描述', '任务未完成的原因', '对后续的影响'],
                                            ['', '', ''], ['', '', ''], ['', '', '']])
    create_table_one_cell(document, '存在的问题及解决方案（必填）')
    create_table_more_cell(document, 5, 4, [['问题描述及原因分析', '解决方案', '预计完成日期', '负责人'],
                                            ['', '', '', ''], ['', '', '', ''], ['', '', '', ''],
                                            ['', '', '', '']])
    create_table_one_cell(document, '说明：如需求、技术方案有变化，请将信的需求文档、技术方案文档与周报一起，提交给公司归档')
    create_table_one_cell(document, '项目进展和计划')
    create_table_one_cell(document, '一、本周工作完成情况（ {0}日至 {1}） （以下必填）'.format(first_date, end_date))
    create_table_more_cell(document, 12, 4, [['编号', '本周重要里程碑事件', '完成日期', '完成标志'],
                                             ['1', '', '', ''], ['2', '', '', ''], ['', '', '', ''],
                                             ['编号', '上周计划的工作内容，但本周已完成', '完成日期', '负责人'],
                                             ['1', '', '', ''], ['2', '', '', ''], ['3', '', '', ''],
                                             ['4', '', '', ''], ['5', '', '', ''], ['6', '', '', ''],
                                             ['7', '', '', '']])


def create_fixed_cell_tow(document):
    """表后半部分固定内容"""
    create_table_one_cell(document, '项目组下周预计借支情况')
    create_table_more_cell(document, 5, 3, [['借支内容摘要', '金额', '备注'], ['', '', ''], ['', '', ''],
                                            ['合计', '', '']])
    create_table_one_cell(document, '已提交给客户的阶段性文档和代码（必填）')
    create_table_more_cell(document, 4, 4, [['资料名称', '提交时间', '接收人', '备注']])
    create_table_one_cell(document, '已提交给公司的阶段性文档和代码（必填）')
    create_table_more_cell(document, 4, 4, [['资料名称', '提交时间', '接收人', '备注']])
    create_table_one_cell(document, '负责人对此项目本周工作的反馈意见')
    create_table_more_cell(document, 3, 2, [['对项目进展评价', ''],
                                            ['对“项目情况”中，变更情况及存在问题的评述', ''],
                                            ['后续项目实施建议', '']])


def to_excel(worksheet, first_date, end_date):
    style = xlwt.XFStyle()
    title_str = '新一代核心系统建设项目周报\n' \
                '\n' \
                '(周期:{0}至{1})'.format(first_date, end_date)
    bl = xlwt.Borders()
    bl.left = xlwt.Borders.THIN
    bl.right = xlwt.Borders.THIN
    bl.top = xlwt.Borders.THIN
    bl.bottom = xlwt.Borders.THIN
    al = xlwt.Alignment()
    al.horz = 0x02  # 设置水平居中
    al.vert = 0x01  # 设置垂直居中
    al.wrap = 1  # 自动换行
    style.alignment = al
    style.borders = bl
    worksheet.write_merge(0, 3, 0, 9, title_str, style)
    worksheet.write_merge(SITE_1, SITE_2-1, 0, 0, '一.本周计划进展情况', style)
    worksheet.write(SITE_1, 1, '序号', style)
    worksheet.write(SITE_1, 2, '工作事项名称', style)
    worksheet.write(SITE_1, 3, '开始时间', style)
    worksheet.write(SITE_1, 4, '完成时间', style)
    worksheet.write(SITE_1, 5, '责任人', style)
    worksheet.write(SITE_1, 6, '计划%', style)
    worksheet.write(SITE_1, 7, '实际%', style)
    worksheet.write(SITE_1, 8, '偏差%', style)
    worksheet.write(SITE_1, 9, '进展说明', style)
    worksheet.write_merge(SITE_2, SITE_2+31, 0, 0, '二.下周工作计划', style)
    worksheet.write(SITE_2, 1, '序号', style)
    worksheet.write(SITE_2, 2, '工作事项名称', style)
    worksheet.write(SITE_2, 3, '开始时间', style)
    worksheet.write(SITE_2, 4, '完成时间', style)
    worksheet.write(SITE_2, 5, '责任人', style)
    worksheet.write_merge(SITE_2, SITE_2, 6, 8, '计划输出结果', style)
    worksheet.write(SITE_2, 9, '说明', style)
    worksheet.write_merge(SITE_2+32, SITE_2+41, 0, 0, '三.目前存在的问题以及需要协调解决的事项', style)
    worksheet.write(SITE_2+32, 1, '序号', style)
    worksheet.write(SITE_2+32, 2, '问题名称', style)
    worksheet.write_merge(SITE_2+32, SITE_2+32, 3, 4, '问题描述', style)
    worksheet.write(SITE_2+32, 5, '提出日期', style)
    worksheet.write(SITE_2+32, 6, '提出人团体', style)
    worksheet.write(SITE_2+32, 7, '解决责任团队', style)
    worksheet.write(SITE_2+32, 8, '预期解决时间', style)
    worksheet.write(SITE_2+32, 9, '解决建议方案和计划', style)
    worksheet.write_merge(SITE_2+42, SITE_2+47, 0, 0, '四.本周质量管理方面的工作总结', style)
    worksheet.write(SITE_2+42, 1, '序号', style)
    worksheet.write_merge(SITE_2+42, SITE_2+42, 2, 9, '进展说明', style)
    worksheet.write_merge(SITE_2+48, SITE_2+53, 0, 0, '五.本周配置管理方面的工作总结', style)
    worksheet.write(SITE_2+48, 1, '序号', style)
    worksheet.write_merge(SITE_2+48, SITE_2+48, 2, 9, '进展说明', style)


def main():
    site_1 = SITE_1
    site_2 = SITE_2
    time_now = date.today()
    # time_now = date(2019, 7, 26)
    today = time_now.strftime("%Y-%m-%d")
    first_date = (time_now + timedelta(days=-4)).strftime("%Y-%m-%d")
    end_date = (time_now + timedelta(days=2)).strftime("%Y-%m-%d")
    first_date_of = (time_now + timedelta(days=3)).strftime("%Y-%m-%d")
    end_date_of = (time_now + timedelta(days=7)).strftime("%Y-%m-%d")
    # 生成excel表格
    workbook = xlwt.Workbook()
    worksheet = workbook.add_sheet('周报', cell_overwrite_ok=True)
    to_excel(worksheet, first_date, end_date)
    # 获取第几周
    week = get_week_of_month(time_now.year, time_now.month, time_now.day)
    # week = get_week_of_month(2019, 8, 2)
    document = Document()
    document.add_heading('项目周报（{0}）'.format(week.strip('#').strip()), level=1)
    document.add_paragraph('填表人：廖虹媛    报告周期：{date1}到{date2}   填表日期：{date3}'.format(
        date1=first_date, date2=end_date, date3=today))
    # # 创建固定列表函数
    # create_fixed_cell(document, first_date, end_date)
    # 本周工作内容表格
    table1 = document.add_table(rows=CELL, cols=6, style='Table Grid')
    table1.rows[0].cells[0].add_paragraph('编号')
    table1.rows[0].cells[1].add_paragraph('本周工作内容')
    table1.rows[0].cells[2].add_paragraph('计划完成时间')
    table1.rows[0].cells[3].add_paragraph('实际完成时间')
    table1.rows[0].cells[4].add_paragraph('负责人')
    table1.rows[0].cells[5].add_paragraph('项目组')
    # 下周工作内容表格
    create_table_one_cell(document, '项目进展和计划')
    create_table_one_cell(document, '一、下周工作完成情况（ {0}至 {1}） （以下必填）'.format(first_date, end_date))
    table2 = document.add_table(rows=CELL, cols=6, style='Table Grid')
    table2.rows[0].cells[0].add_paragraph('编号')
    table2.rows[0].cells[1].add_paragraph('下周工作内容')
    table2.rows[0].cells[2].add_paragraph('计划完成时间')
    table2.rows[0].cells[3].add_paragraph('实际完成时间')
    table2.rows[0].cells[4].add_paragraph('负责人')
    table2.rows[0].cells[5].add_paragraph('项目组')
    # 主要内容写入
    get_file_path(MD_PATH, week, table1, table2, first_date, today, worksheet,
                  site_1, site_2, first_date_of, end_date_of)
    # # 后半部函数
    # create_fixed_cell_tow(document)
    save_name = '厦开项目组周报{0}至{1}.docx'.format(first_date, end_date)
    document.save(save_name)
    excel_name = '新一代核心系统建设项目周报{0}_天用厦开安全项目组.xls'.format(end_date)
    workbook.save(excel_name)


if __name__ == '__main__':
    main()
